import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
import random
import string
import re
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import pywhatkit as ws
from docx import Document
from email_validator import validate_email, EmailNotValidError
import phonenumbers
from tabulate import tabulate
from datetime import datetime

clientes = []
productos = []
ventas = []
ventas_registradas = []

SMTP_SERVER = 'smtp.gmail.com'
SMTP_PORT = 587
EMAIL_ADDRESS = os.getenv('EMAIL_ADDRESS')
EMAIL_PASSWORD = os.getenv('EMAIL_PASSWORD')

def enviar_correo_bienvenida(correo, nombre, apellido, codigo_cliente):
    mensaje = MIMEMultipart()
    mensaje['From'] = EMAIL_ADDRESS
    mensaje['To'] = correo
    mensaje['Subject'] = "Bienvenido a nuestra clientela"
    
    body = f"Hola {nombre} {apellido},\n\nBienvenido a nuestra clientela. Gracias por unirse. Su código de cliente es {codigo_cliente}."
    mensaje.attach(MIMEText(body, 'plain'))
    
    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        server.sendmail(EMAIL_ADDRESS, correo, mensaje.as_string())
        server.quit()
    except Exception as e:
        print(f"Error al enviar el correo: {e}")

    with open("PyWhatKit_DB.txt", "a") as f:
        f.write(f"Correo enviado a: {correo} con el código de cliente: {codigo_cliente}\n")

def enviar_correo_edicion(correo, nombre, apellido):
    mensaje = MIMEMultipart()
    mensaje['From'] = EMAIL_ADDRESS
    mensaje['To'] = correo
    mensaje['Subject'] = "Actualización de Información"
    
    body = f"Hola {nombre} {apellido},\n\nSu información ha sido actualizada exitosamente en nuestro sistema."
    mensaje.attach(MIMEText(body, 'plain'))
    
    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        server.sendmail(EMAIL_ADDRESS, correo, mensaje.as_string())
        server.quit()
    except Exception as e:
        print(f"Error al enviar el correo: {e}")

    with open("PyWhatKit_DB.txt", "a") as f:
        f.write(f"Correo enviado a: {correo} notificando actualización.\n")

def enviar_mensaje_whatsapp(numero, nombre, apellido, codigo_cliente=None, tipo=None):
    if tipo == "edicion":
        mensaje = f"Hola {nombre} {apellido}, tu información ha sido actualizada exitosamente en nuestro sistema."
    else:
        mensaje = f"Hola {nombre} {apellido}, tu código de cliente es {codigo_cliente}."

    ws.sendwhatmsg_instantly(numero, mensaje)
    with open("PyWhatKit_DB.txt", "a") as f:
        f.write(f"WhatsApp enviado a: {numero} con el mensaje: {mensaje}\n")

def crear_ventana_inicio():
    root = tk.Tk()
    root.title("Inicio de Sesión")
    root.geometry("600x500")
    root.config(bg="#212020")

    frame = ttk.Frame(root, padding=20)
    frame.pack(expand=True)

    ttk.Label(frame, text="Ingrese el correo para enviar notificaciones:", background="#212020", foreground="white").pack(pady=5)
    global email_entry
    email_entry = ttk.Entry(frame, width=40)
    email_entry.pack(pady=5)

    ttk.Label(frame, text="Ingrese la contraseña del correo:", background="#212020", foreground="white").pack(pady=5)
    global password_entry
    password_entry = ttk.Entry(frame, show="*", width=40)
    password_entry.pack(pady=5)

    ttk.Button(frame, text="Guardar y Continuar", command=guardar_configuracion).pack(pady=10)

    return root

def guardar_configuracion():
    global EMAIL_ADDRESS, EMAIL_PASSWORD
    EMAIL_ADDRESS = email_entry.get()
    EMAIL_PASSWORD = password_entry.get()
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        messagebox.showerror("Error", "Por favor, complete ambos campos.")
        return
    
    # Ocultar ventana de inicio de sesión
    root.withdraw()
    
    # Mostrar bienvenida
    bienvenida = tk.Toplevel()
    bienvenida.title("Bienvenida")
    bienvenida.geometry("600x500")
    bienvenida.config(bg="#212020")
    
    ttk.Label(bienvenida, text="¡Bienvenido al Sistema!", font=('Arial', 14), background="#212020", foreground="white").pack(pady=50)
    
    # Cerrar la ventana de bienvenida después de 3 segundos
    root.after(3000, lambda: (bienvenida.destroy(), mostrar_interfaz_principal()))

#-----------------------------------------------------------------VENTANA PRINCIPAL
    
def mostrar_interfaz_principal():
    global root
    root = tk.Tk()
    root.title("Sistema de Gestión")
    root.state('zoomed')
    root.config(bg="#212020")

    label_bienvenida = ttk.Label(root, text="Bienvenido al Sistema", font=('Arial', 14), background="#212020", foreground="white")
    label_bienvenida.pack(pady=20)

    button_frame = ttk.Frame(root)
    button_frame.pack(pady=20, padx=20, fill=tk.BOTH, expand=True)

    button_grid_frame = ttk.Frame(button_frame)
    button_grid_frame.pack(padx=20, pady=20)

    botones = [
        ("Registrar Cliente", mostrar_registrar_cliente),
        ("Consultar Clientes", mostrar_consultar_clientes),
        ("Registrar Producto", mostrar_registrar_producto),
        ("Consultar Productos", mostrar_consultar_productos),
        ("Registrar Venta", mostrar_registrar_venta),
        ("Consultar Ventas", mostrar_consultar_ventas)
    ]

    for i, (texto, comando) in enumerate(botones):
        button = ttk.Button(button_grid_frame, text=texto, command=lambda func=comando: mostrar_ventana(func))
        button.grid(row=i//2, column=i%2, padx=10, pady=10, sticky='nsew')

    button_salir = ttk.Button(root, text="Salir del Sistema", command=root.quit)
    button_salir.pack(pady=10)

    button_grid_frame.grid_rowconfigure(0, weight=1)
    button_grid_frame.grid_rowconfigure(1, weight=1)
    button_grid_frame.grid_rowconfigure(2, weight=1)
    button_grid_frame.grid_columnconfigure(0, weight=1)
    button_grid_frame.grid_columnconfigure(1, weight=1)

    root.mainloop()

def mostrar_ventana(funcion):
    global root
    root.withdraw()
    
    ventana = tk.Toplevel()
    ventana.title(funcion.__name__)
    ventana.state('zoomed')
    ventana.config(bg="#212020")
    ventana.pack_propagate(False)

    button_volver = ttk.Button(ventana, text="Volver a la pantalla principal", command=lambda: volver_a_principal(ventana))
    button_volver.pack(pady=10)

    funcion(ventana)

def volver_a_principal(ventana):
    ventana.destroy()
    root.deiconify()

def generar_codigo_cliente():
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))

def validar_campos(datos_cliente):
    for campo, valor in datos_cliente.items():
        if not valor:
            messagebox.showerror("Error de validación", f"El campo '{campo}' no puede estar vacío.")
            return False

    try:
        validate_email(datos_cliente["Correo"])
    except EmailNotValidError:
        messagebox.showerror("Error de validación", "El correo electrónico no es válido.")
        return False

    try:
        numero = phonenumbers.parse(datos_cliente["Teléfono"], None)
        if not phonenumbers.is_valid_number(numero):
            raise ValueError
    except (phonenumbers.NumberParseException, ValueError):
        messagebox.showerror("Error de validación", "El número de teléfono no es válido.")
        return False

    if not re.match(r"^\d{2}/\d{2}/\d{2}$", datos_cliente["Fecha de nacimiento"]):
        messagebox.showerror("Error de validación", "La fecha de nacimiento debe estar en el formato dd/mm/aa.")
        return False

    return True

#------------------------------------------------------VENTANA DE REGISTRO DE CLIENTES

def mostrar_registrar_cliente(ventana):
    tk.Label(ventana, text="Registro de Cliente", font=('Arial', 12), background="#212020", foreground="white").pack(pady=10)

    # Crear un frame para centrar el formulario
    form_frame = ttk.Frame(ventana, padding=10)
    form_frame.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

    campos = ["Nombre", "Apellido", "Correo", "Teléfono", "Dirección", "Ciudad", "País", "Fecha de nacimiento", "Género"]
    entradas = {}

    for campo in campos:
        label = ttk.Label(form_frame, text=campo + ":")
        label.pack(anchor=tk.W, pady=2)
        entrada = ttk.Entry(form_frame, width=30)
        entrada.pack(anchor=tk.W, pady=2)
        entradas[campo] = entrada

    codigo_cliente = generar_codigo_cliente()
    ttk.Label(form_frame, text="Código de cliente:").pack(anchor=tk.W, pady=2)
    ttk.Label(form_frame, text=codigo_cliente).pack(anchor=tk.W, pady=2)

    button_registrar = ttk.Button(form_frame, text="Registrar Cliente", command=lambda: registrar_cliente(ventana, entradas, codigo_cliente))
    button_registrar.pack(pady=10)

def registrar_cliente(ventana, entradas, codigo_cliente):
    global clientes
    datos_cliente = {campo: entrada.get() for campo, entrada in entradas.items()}
    datos_cliente["Código de cliente"] = codigo_cliente

    if validar_campos(datos_cliente):
        clientes.append(datos_cliente)
        
        with open("clientes_exportados.txt", "a") as f:
            f.write(f"Cliente registrado: {datos_cliente}\n")

        # Preguntar cómo notificar al cliente
        notificar = simpledialog.askstring("Notificación", "¿Cómo deseas notificar al cliente?\n1. Correo\n2. WhatsApp\n3. Ambos\nIngresa 1, 2, o 3:")
        
        if notificar in ["1", "3"]:
            # Enviar correo de bienvenida
            enviar_correo_bienvenida(datos_cliente["Correo"], datos_cliente["Nombre"], datos_cliente["Apellido"], codigo_cliente)
        
        if notificar in ["2", "3"]:
            # Enviar mensaje de WhatsApp
            enviar_mensaje_whatsapp(datos_cliente["Teléfono"], datos_cliente["Nombre"], datos_cliente["Apellido"], codigo_cliente)

        messagebox.showinfo("Registro exitoso", "Cliente registrado exitosamente.")
        ventana.destroy()
        root.deiconify()

def mostrar_consultar_clientes(ventana):
    ventana.title("Consultar Clientes")

    columns = ["Código de cliente", "Nombre", "Apellido", "Correo", "Teléfono", "Dirección", "Ciudad", "País", "Fecha de nacimiento", "Género"]
    
    # Crear un frame principal para contener todo el contenido
    main_frame = ttk.Frame(ventana)
    main_frame.pack(pady=10, fill=tk.BOTH, expand=True)

    # Frame para el Treeview y la barra de desplazamiento
    tree_frame = ttk.Frame(main_frame)
    tree_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    tree = ttk.Treeview(tree_frame, columns=columns, show='headings')
    tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Añadir barra de desplazamiento vertical
    scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    tree.configure(yscroll=scrollbar.set)

    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150, anchor='center')

    cargar_datos_clientes(tree)

    # Frame para el campo de búsqueda y el botón de exportar
    search_export_frame = ttk.Frame(main_frame)
    search_export_frame.pack(side=tk.TOP, pady=10, fill=tk.X)

    # Campo de búsqueda
    label_buscar = ttk.Label(search_export_frame, text="Buscar por código de cliente:")
    label_buscar.pack(side=tk.LEFT, padx=5)

    entrada_buscar = ttk.Entry(search_export_frame, width=30)
    entrada_buscar.pack(side=tk.LEFT, padx=5)
    entrada_buscar.bind('<KeyRelease>', lambda event: buscar_cliente_dinamico(entrada_buscar.get(), tree))

    # Botón de exportar clientes a Word
    button_exportar = ttk.Button(search_export_frame, text="Exportar Clientes a Word", command=lambda: exportar_clientes_word(tree))
    button_exportar.pack(side=tk.LEFT, padx=5)

    # Frame para los botones de acción
    action_button_frame = ttk.Frame(main_frame)
    action_button_frame.pack(side=tk.BOTTOM, pady=10, fill=tk.X)

    button_editar = ttk.Button(action_button_frame, text="Editar Cliente", command=lambda: editar_cliente(tree))
    button_editar.pack(side=tk.LEFT, padx=5)

    button_eliminar = ttk.Button(action_button_frame, text="Eliminar Cliente", command=lambda: eliminar_cliente(tree))
    button_eliminar.pack(side=tk.LEFT, padx=5)

def buscar_cliente_dinamico(busqueda, tree):
    busqueda = busqueda.lower()
    for item in tree.get_children():
        tree.delete(item)
    for cliente in clientes:
        if busqueda in cliente["Código de cliente"].lower():
            tree.insert("", tk.END, values=[cliente.get(col, "") for col in tree.cget("columns")])

def cargar_datos_clientes(tree):
    for cliente in clientes:
        tree.insert("", tk.END, values=[cliente.get(col, "") for col in tree.cget("columns")])

def editar_cliente(tree):
    seleccionado = tree.focus()
    if not seleccionado:
        messagebox.showwarning("Advertencia", "Seleccione un cliente para editar.")
        return

    datos_cliente = tree.item(seleccionado, "values")
    campos = ["Código de cliente", "Nombre", "Apellido", "Correo", "Teléfono", "Dirección", "Ciudad", "País", "Fecha de nacimiento", "Género"]

    ventana_editar = tk.Toplevel()
    ventana_editar.title("Editar Cliente")
    ventana_editar.geometry("600x500")
    ventana_editar.config(bg="#212020")

    form_frame = ttk.Frame(ventana_editar, padding=20)
    form_frame.pack(pady=10, fill=tk.BOTH, expand=True)

    entradas = {}
    for i, campo in enumerate(campos):
        label = ttk.Label(form_frame, text=campo + ":", background="#212020", foreground="white")
        label.grid(row=i, column=0, sticky=tk.W, pady=5)
        entrada = ttk.Entry(form_frame, width=50)
        entrada.grid(row=i, column=1, pady=5)
        if campo == "Código de cliente":
            entrada.insert(0, datos_cliente[i])
            entrada.config(state='readonly')  # Código no editable
        else:
            entrada.insert(0, datos_cliente[i])
        entradas[campo] = entrada

    button_actualizar = ttk.Button(form_frame, text="Actualizar Cliente", command=lambda: actualizar_cliente(tree, seleccionado, entradas, ventana_editar))
    button_actualizar.grid(row=len(campos), column=0, columnspan=2, pady=20)

def actualizar_cliente(tree, item, entradas, ventana_editar):
    global clientes
    datos_actualizados = {campo: entrada.get() for campo, entrada in entradas.items()}
    
    if not messagebox.askyesno("Confirmación", "¿Desea notificar al cliente sobre la actualización?"):
        notificar = "0"
    else:
        notificar = simpledialog.askstring("Notificación", "¿Cómo deseas notificar al cliente?\n1. Correo\n2. WhatsApp\n3. Ambos\nIngresa 1, 2, o 3:")
    
    for i, cliente in enumerate(clientes):
        if cliente["Código de cliente"] == datos_actualizados["Código de cliente"]:
            clientes[i] = datos_actualizados
            break

    tree.item(item, values=[datos_actualizados.get(col, "") for col in tree.cget("columns")])
    ventana_editar.destroy()
    
    if notificar in ["1", "3"]:
        enviar_correo_edicion(datos_actualizados["Correo"], datos_actualizados["Nombre"], datos_actualizados["Apellido"])

    if notificar in ["2", "3"]:
        enviar_mensaje_whatsapp(datos_actualizados["Teléfono"], datos_actualizados["Nombre"], datos_actualizados["Apellido"], tipo="edicion")

def eliminar_cliente(tree):
    global clientes
    seleccionado = tree.focus()
    if not seleccionado:
        messagebox.showwarning("Advertencia", "Seleccione un cliente para eliminar.")
        return

    if not messagebox.askyesno("Confirmación", "¿Está seguro de que desea eliminar este cliente?"):
        return

    datos_cliente = tree.item(seleccionado, "values")
    codigo_cliente = datos_cliente[0]

    clientes = [cliente for cliente in clientes if cliente["Código de cliente"] != codigo_cliente]
    tree.delete(seleccionado)

def exportar_clientes_word(tree):
    doc = Document()
    doc.add_heading('Clientes Registrados', 0)

    table = doc.add_table(rows=1, cols=len(tree.cget("columns")))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(tree.cget("columns")):
        hdr_cells[i].text = col

    for item in tree.get_children():
        row = table.add_row().cells
        for i, value in enumerate(tree.item(item, "values")):
            row[i].text = str(value)

    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Exportación exitosa", "Clientes exportados a Word exitosamente.")

def generar_codigo_cliente():
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))

def validar_campos(datos_cliente):
    for campo, valor in datos_cliente.items():
        if not valor:
            messagebox.showerror("Error", f"El campo '{campo}' es obligatorio.")
            return False
    return True

#---------------------------------------------------------------REGISTRAR PRODUCTO 


def mostrar_registrar_producto(ventana):
    tk.Label(ventana, text="Registro de Producto", font=('Arial', 12), background="#212020", foreground="white").pack(pady=10)

    form_frame = ttk.Frame(ventana, padding=10)
    form_frame.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

    campos = ["SKU", "Nombre", "Descripción", "Precio", "Stock", "Categoría", "Marca", "Fabricante"]
    entradas = {}

    for campo in campos:
        label = ttk.Label(form_frame, text=campo + ":")
        label.pack(anchor=tk.W, pady=2)
        entrada = ttk.Entry(form_frame, width=30)
        entrada.pack(anchor=tk.W, pady=2)
        entradas[campo] = entrada

    button_registrar = ttk.Button(form_frame, text="Registrar Producto", command=lambda: registrar_producto(ventana, entradas))
    button_registrar.pack(pady=10)

def validar_campos_producto(datos_producto):
    if not datos_producto["SKU"].isdigit():
        messagebox.showerror("Error de validación", "El SKU debe ser un código único numérico.")
        return False

    if not re.match(r"^\d+(\.\d{2})?$", datos_producto["Precio"]):
        messagebox.showerror("Error de validación", "El precio debe estar en el formato 00.00.")
        return False

    try:
        int(datos_producto["Stock"])
    except ValueError:
        messagebox.showerror("Error de validación", "El stock debe ser un valor entero.")
        return False

    return True

def registrar_producto(ventana, entradas):
    global productos

    datos_producto = {campo: entrada.get() for campo, entrada in entradas.items()}

    if not validar_campos_producto(datos_producto):
        return

    productos.append(datos_producto)

    with open("productos_exportados.txt", "a") as f:
        f.write(f"Producto registrado: {datos_producto}\n")

    messagebox.showinfo("Registro exitoso", "Producto registrado exitosamente.")

    respuesta = messagebox.askyesno("Agregar Otro Producto", "¿Deseas agregar otro producto?")
    
    if respuesta:
        for entrada in entradas.values():
            entrada.delete(0, tk.END)
    else:
        ventana.destroy()
        root.deiconify()

def mostrar_consultar_productos(ventana):
    ventana.title("Consultar Productos")

    columns = ["SKU", "Nombre", "Descripción", "Precio", "Stock", "Categoría", "Marca", "Fabricante"]

    main_frame = ttk.Frame(ventana)
    main_frame.pack(pady=10, fill=tk.BOTH, expand=True)

    tree_frame = ttk.Frame(main_frame)
    tree_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    tree = ttk.Treeview(tree_frame, columns=columns, show='headings')
    tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    tree.configure(yscroll=scrollbar.set)

    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150, anchor='center')

    cargar_datos_productos(tree)

    search_export_frame = ttk.Frame(main_frame)
    search_export_frame.pack(side=tk.TOP, pady=10, fill=tk.X)

    label_buscar = ttk.Label(search_export_frame, text="Buscar por SKU:")
    label_buscar.pack(side=tk.LEFT, padx=5)

    entrada_buscar = ttk.Entry(search_export_frame, width=30)
    entrada_buscar.pack(side=tk.LEFT, padx=5)
    entrada_buscar.bind('<KeyRelease>', lambda event: buscar_producto_dinamico(entrada_buscar.get(), tree))

    button_exportar = ttk.Button(search_export_frame, text="Exportar Productos a Word", command=lambda: exportar_productos_word(tree))
    button_exportar.pack(side=tk.LEFT, padx=5)

    action_button_frame = ttk.Frame(main_frame)
    action_button_frame.pack(side=tk.BOTTOM, pady=10, fill=tk.X)

    button_editar = ttk.Button(action_button_frame, text="Editar Producto", command=lambda: editar_producto(tree))
    button_editar.pack(side=tk.LEFT, padx=5)

    button_eliminar = ttk.Button(action_button_frame, text="Eliminar Producto", command=lambda: eliminar_producto(tree))
    button_eliminar.pack(side=tk.LEFT, padx=5)

def cargar_datos_productos(tree):
    for item in tree.get_children():
        tree.delete(item)
    for producto in productos:
        tree.insert("", tk.END, iid=producto["SKU"], values=[producto["SKU"], producto["Nombre"], producto["Descripción"], producto["Precio"], producto["Stock"], producto["Categoría"], producto["Marca"], producto["Fabricante"]])

def buscar_producto_dinamico(busqueda, tree):
    busqueda = busqueda.lower()
    for item in tree.get_children():
        tree.delete(item)
    for producto in productos:
        if busqueda in producto["SKU"].lower():
            tree.insert("", tk.END, iid=producto["SKU"], values=[producto["SKU"], producto["Nombre"], producto["Descripción"], producto["Precio"], producto["Stock"], producto["Categoría"], producto["Marca"], producto["Fabricante"]])

def exportar_productos_word(tree):
    doc = Document()
    doc.add_heading('Productos Registrados', 0)

    columns = tree.cget("columns")
    table = doc.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col

    for item in tree.get_children():
        row = table.add_row().cells
        for i, val in enumerate(tree.item(item, "values")):
            row[i].text = str(val)

    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos de Word", "*.docx")])
    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Exportación exitosa", f"Productos exportados exitosamente a {file_path}")

    with open("productos_exportados.txt", "a") as f:
        f.write(tabulate([tree.item(item, "values") for item in tree.get_children()], headers=columns, tablefmt="plain") + "\n")

def editar_producto(tree):
    seleccionado = tree.focus()
    if not seleccionado:
        messagebox.showwarning("Advertencia", "Seleccione un producto para editar.")
        return

    datos_producto = tree.item(seleccionado, "values")
    campos = ["SKU", "Nombre", "Descripción", "Precio", "Stock", "Categoría", "Marca", "Fabricante"]

    ventana_editar = tk.Toplevel()
    ventana_editar.title("Editar Producto")
    ventana_editar.geometry("600x500")
    ventana_editar.config(bg="#212020")

    form_frame = ttk.Frame(ventana_editar, padding=20)
    form_frame.pack(pady=10, fill=tk.BOTH, expand=True)

    entradas = {}
    for i, campo in enumerate(campos):
        label = ttk.Label(form_frame, text=campo + ":", background="#212020", foreground="white")
        label.grid(row=i, column=0, sticky=tk.W, pady=5)
        entrada = ttk.Entry(form_frame, width=50)
        entrada.grid(row=i, column=1, pady=5)
        if campo == "SKU":
            entrada.insert(0, datos_producto[i])
            entrada.config(state='readonly')
        else:
            entrada.insert(0, datos_producto[i])
        entradas[campo] = entrada

    button_actualizar = ttk.Button(form_frame, text="Actualizar Producto", command=lambda: actualizar_producto(tree, seleccionado, entradas, ventana_editar))
    button_actualizar.grid(row=len(campos), column=0, columnspan=2, pady=20)

def actualizar_producto(tree, item, entradas, ventana_editar):
    global productos

    datos_actualizados = {campo: entrada.get() for campo, entrada in entradas.items()}

    if not validar_campos_producto(datos_actualizados):
        return

    for i, producto in enumerate(productos):
        if producto["SKU"] == datos_actualizados["SKU"]:
            productos[i] = datos_actualizados
            break

    tree.item(item, values=[datos_actualizados.get(col, "") for col in tree.cget("columns")])
    ventana_editar.destroy()

def eliminar_producto(tree):
    global productos

    seleccionado = tree.focus()
    if not seleccionado:
        messagebox.showwarning("Advertencia", "Seleccione un producto para eliminar.")
        return

    if not messagebox.askyesno("Confirmación", "¿Está seguro de que desea eliminar este producto?"):
        return

    datos_producto = tree.item(seleccionado, "values")
    sku_producto = datos_producto[0]

    productos = [producto for producto in productos if producto["SKU"] != sku_producto]

    tree.delete(seleccionado)

#---------------------------------------------------------------REGISTAR VENTA


def calcular_total():
    total = 0
    for item in listbox_productos.get_children():
        valores = listbox_productos.item(item, 'values')
        total_producto = valores[3].replace('$', '')
        total += float(total_producto)
    label_total.config(text=f"Total: ${total:.2f}")

def actualizar_precio(event):
    producto = combo_producto.get()
    precio = next(((p["Precio"]) for p in productos if p["Nombre"] == producto), 0)
    entry_precio.config(state=tk.NORMAL)
    entry_precio.delete(0, tk.END)
    entry_precio.insert(0, f"{float(precio):.2f}")
    entry_precio.config(state=tk.DISABLED)

def agregar_producto():
    producto_nombre = combo_producto.get()
    precio = entry_precio.get()
    cantidad = entry_cantidad.get()

    if producto_nombre and precio and cantidad:
        try:
            cantidad = int(cantidad)  # Asegúrate de que la cantidad sea un entero
        except ValueError:
            messagebox.showwarning("Error de cantidad", "La cantidad debe ser un número entero.")
            return

        # Verificar el stock
        stock_disponible = next((int(p["Stock"]) for p in productos if p["Nombre"] == producto_nombre), 0)
        
        if cantidad > stock_disponible:
            messagebox.showwarning("Stock Insuficiente", f"No hay suficiente stock para {producto_nombre}. Stock disponible: {stock_disponible}.")
            return

        # Agregar producto al Treeview
        total_producto = float(precio) * cantidad
        listbox_productos.insert('', tk.END, values=(
            producto_nombre, precio, cantidad, f"${total_producto:.2f}"
        ))

        # Actualizar stock en la tabla de productos
        for producto in productos:
            if producto["Nombre"] == producto_nombre:
                producto["Stock"] = stock_disponible - cantidad
                break

        # Limpiar campos de entrada
        entry_cantidad.delete(0, tk.END)
        
        # Recalcular el total
        calcular_total()

def registrar_venta():
    cliente_seleccionado = combo_cliente.get()
    productos_vendidos = listbox_productos.get_children()
    total_venta = label_total.cget("text").split('$')[1]
    
    if not cliente_seleccionado or not productos_vendidos:
        messagebox.showwarning("Advertencia", "Debe seleccionar un cliente y agregar al menos un producto.")
        return

    cliente = next((c for c in clientes if f"{c['Nombre']} {c['Apellido']}" == cliente_seleccionado), None)
    codigo_cliente = cliente.get('Código de cliente', "Desconocido") if cliente else "Desconocido"
    
    productos = []
    for item in productos_vendidos:
        valores = listbox_productos.item(item, 'values')
        productos.append({
            "Nombre": valores[0],
            "Cantidad": valores[2],
            "Precio": valores[1],
        })
    
    venta = {
        "ID": len(ventas) + 1,
        "CodigoCliente": codigo_cliente,
        "Nombre": cliente['Nombre'] if cliente else "Desconocido",
        "Apellido": cliente['Apellido'] if cliente else "Desconocido",
        "Productos": productos,
        "Total": float(total_venta),
        "FechaHora": datetime.now()
    }
    
    ventas.append(venta)

    messagebox.showinfo("Venta Registrada", f"Venta registrada con éxito.\n\nCliente: {cliente_seleccionado}\nTotal: ${total_venta}")

    combo_cliente.set('')
    listbox_productos.delete(*listbox_productos.get_children())
    label_total.config(text="Total: $0.00")

def mostrar_registrar_venta(ventana):
    ventana.title("Sistema de Caja")
    ventana.config(bg="#2b2b2b")

    # Marco para la información del cliente
    frame_cliente = tk.Frame(ventana, bg="#2b2b2b")
    frame_cliente.pack(padx=10, pady=10, fill=tk.X)
    
    tk.Label(frame_cliente, text="Cliente:", bg="#2b2b2b", fg="white").pack(side=tk.LEFT)
    global combo_cliente
    combo_cliente = ttk.Combobox(frame_cliente, values=[f"{cliente['Nombre']} {cliente['Apellido']}" for cliente in clientes])
    combo_cliente.pack(side=tk.LEFT, padx=5)

    # Marco para los productos
    frame_productos = tk.Frame(ventana, bg="#2b2b2b")
    frame_productos.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

    tk.Label(frame_productos, text="Producto:", bg="#2b2b2b", fg="white").pack(side=tk.LEFT)
    
    global combo_producto
    combo_producto = ttk.Combobox(frame_productos)
    combo_producto['values'] = [producto["Nombre"] for producto in productos]
    combo_producto.pack(side=tk.LEFT, padx=5)
    combo_producto.bind("<<ComboboxSelected>>", actualizar_precio)
    combo_producto.bind('<KeyRelease>', filtrar_producto)
    
    tk.Label(frame_productos, text="Precio:", bg="#2b2b2b", fg="white").pack(side=tk.LEFT)
    global entry_precio
    entry_precio = tk.Entry(frame_productos, state=tk.DISABLED)
    entry_precio.pack(side=tk.LEFT, padx=5)
    
    tk.Label(frame_productos, text="Cantidad:", bg="#2b2b2b", fg="white").pack(side=tk.LEFT)
    global entry_cantidad
    entry_cantidad = tk.Entry(frame_productos)
    entry_cantidad.pack(side=tk.LEFT, padx=5)
    
    tk.Button(frame_productos, text="Agregar Producto", command=agregar_producto, bg="#4caf50", fg="white").pack(side=tk.LEFT, padx=5)
    
    global listbox_productos
    listbox_productos = ttk.Treeview(frame_productos, columns=("Nombre", "Precio", "Cantidad", "Total"), show='headings')
    listbox_productos.heading("Nombre", text="Nombre")
    listbox_productos.heading("Precio", text="Precio")
    listbox_productos.heading("Cantidad", text="Cantidad")
    listbox_productos.heading("Total", text="Total")
    listbox_productos.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)
    
    # Marco para el total
    frame_total = tk.Frame(ventana, bg="#2b2b2b")
    frame_total.pack(padx=10, pady=10, fill=tk.X)
    
    global label_total
    label_total = tk.Label(frame_total, text="Total: $0.00", bg="#2b2b2b", fg="white", font=("Helvetica", 24, "bold"))
    label_total.pack(side=tk.LEFT)
    
    # Botón para registrar la venta
    tk.Button(frame_total, text="Registrar Venta", command=registrar_venta, bg="#4CAF50", fg="white", font=("Helvetica", 16, "bold"), relief=tk.RAISED, bd=4).pack(side=tk.RIGHT, padx=5)


def filtrar_producto(event):
    valor = combo_producto.get().lower()
    productos_filtrados = [producto["Nombre"] for producto in productos if valor in producto["Nombre"].lower()]
    combo_producto['values'] = productos_filtrados
    combo_producto.event_generate('<Down>')

def eliminar_venta():
    selected_item = tree.selection()
    if not selected_item:
        messagebox.showwarning("Selección requerida", "Por favor, selecciona una venta para eliminar.")
        return

    venta_id = selected_item[0]  # Obtener el ID del ítem seleccionado

    if messagebox.askyesno("Confirmar Eliminación", f"¿Estás seguro de que deseas eliminar la venta con ID {venta_id}?"):
        global ventas
        ventas = [venta for venta in ventas if venta["ID"] != int(venta_id)]
        cargar_datos_ventas(tree)

def mostrar_consultar_ventas(ventana):
    ventana.title("Consultar Ventas")

    columns = ["Código Cliente", "Nombre", "Apellido", "Productos", "Total", "Fecha y Hora"]

    main_frame = ttk.Frame(ventana)
    main_frame.pack(pady=10, fill=tk.BOTH, expand=True)

    tree_frame = ttk.Frame(main_frame)
    tree_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    global tree
    tree = ttk.Treeview(tree_frame, columns=columns, show='headings')
    tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    tree.configure(yscroll=scrollbar.set)

    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150, anchor='center')

    cargar_datos_ventas(tree)

    search_export_frame = ttk.Frame(main_frame)
    search_export_frame.pack(side=tk.TOP, pady=10, fill=tk.X)

    label_buscar = ttk.Label(search_export_frame, text="Buscar por código de cliente:")
    label_buscar.pack(side=tk.LEFT, padx=5)

    entrada_buscar = ttk.Entry(search_export_frame, width=30)
    entrada_buscar.pack(side=tk.LEFT, padx=5)
    entrada_buscar.bind('<KeyRelease>', lambda event: buscar_venta_dinamica(entrada_buscar.get(), tree))

    button_exportar = ttk.Button(search_export_frame, text="Exportar Ventas a Word", command=lambda: exportar_ventas_word(tree))
    button_exportar.pack(side=tk.LEFT, padx=5)

    action_button_frame = ttk.Frame(main_frame)
    action_button_frame.pack(side=tk.BOTTOM, pady=10, fill=tk.X)

    button_eliminar = ttk.Button(action_button_frame, text="Eliminar Venta", command=eliminar_venta)
    button_eliminar.pack(side=tk.LEFT, padx=5)

    button_ver_detalles = ttk.Button(action_button_frame, text="Ver Detalles", command=ver_detalles_venta)
    button_ver_detalles.pack(side=tk.LEFT, padx=5)


def cargar_datos_ventas(tree):
    for item in tree.get_children():
        tree.delete(item)
    for venta in ventas:
        tree.insert("", tk.END, iid=venta["ID"], values=[
            venta["CodigoCliente"],
            venta["Nombre"],
            venta["Apellido"],
            "**********",  # Muestra "**********" en lugar de los productos
            f"${venta['Total']:.2f}",
            venta["FechaHora"].strftime("%Y-%m-%d %H:%M:%S")
        ])

def ver_detalles_venta():
    selected_item = tree.selection()
    if not selected_item:
        messagebox.showwarning("Selección requerida", "Por favor, selecciona una venta para ver los detalles.")
        return

    venta_id = selected_item[0]  # Obtener el ID del ítem seleccionado
    venta = next((v for v in ventas if v["ID"] == int(venta_id)), None)
    if venta:
        detalles_venta = "\n".join([f"Producto: {p['Nombre']}\nCantidad: {p['Cantidad']}\nPrecio: {p['Precio']}\nTotal: {float(p['Precio'].replace('$', '')) * int(p['Cantidad']):.2f}" for p in venta['Productos']])
        messagebox.showinfo("Detalles de Venta", f"Detalles de la Venta ID {venta_id}:\n\n{detalles_venta}")
    else:
        messagebox.showwarning("Error", "No se encontraron detalles para esta venta.")

def buscar_venta_dinamica(busqueda, tree):
    busqueda = busqueda.lower()
    for item in tree.get_children():
        tree.delete(item)
    for venta in ventas:
        if busqueda in venta["CodigoCliente"].lower():
            productos = "\n".join([f"{p['Nombre']} - Cantidad: {p['Cantidad']} - Precio: {p['Precio']}" for p in venta['Productos']])
            tree.insert("", tk.END, values=[
                venta["CodigoCliente"],
                venta["Nombre"],
                venta["Apellido"],
                productos,
                f"${venta['Total']:.2f}",
                venta["FechaHora"].strftime("%Y-%m-%d %H:%M:%S")
            ])

def exportar_ventas_word(tree):
    doc = Document()
    doc.add_heading('Ventas Registradas', 0)

    columns = tree.cget("columns")
    table = doc.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col

    for item in tree.get_children():
        row = table.add_row().cells
        for i, val in enumerate(tree.item(item, "values")):
            row[i].text = str(val)

    try:
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos de Word", "*.docx")])
        if file_path:
            doc.save(file_path)
            messagebox.showinfo("Exportación exitosa", f"Ventas exportadas exitosamente a {file_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Error al exportar a Word: {e}")

    try:
        with open("ventas_exportadas.txt", "a") as f:
            f.write(tabulate([tree.item(item, "values") for item in tree.get_children()], headers=columns, tablefmt="plain") + "\n")
    except Exception as e:
        messagebox.showerror("Error", f"Error al guardar archivo de texto: {e}")

if __name__ == "__main__":

    root = crear_ventana_inicio()
    root.mainloop()